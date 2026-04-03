from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt


ROOT = Path(__file__).resolve().parents[1]
AI_DIR = ROOT / "doc" / "ai"
HUMAN_DIR = ROOT / "doc" / "human"
SKIP_FILES = {"README.md", "HUMAN_README.md"}
GENERATED_ON = date.today().isoformat()

INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|__.+?__|`.+?`|\*[^*\n][^*\n]*\*)", re.DOTALL)
HORIZONTAL_RULE_PATTERN = re.compile(r"^[-*_]{3,}$")
ORDER_LINE_PATTERN = re.compile(r"^\d+\.\s+`(?P<name>[^`]+)`\s*$")


def set_run_font(run, font_name: str, size_pt: int | None = None, bold: bool | None = None):
    run.font.name = font_name
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


def configure_style(doc: Document, style_name: str, font_name: str, size_pt: int, *, bold: bool = False):
    style = doc.styles[style_name]
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15


def configure_document(doc: Document, title: str):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "TH Sarabun New"
    normal.font.size = Pt(16)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    configure_style(doc, "Title", "TH Sarabun New", 26, bold=True)
    configure_style(doc, "Heading 1", "TH Sarabun New", 20, bold=True)
    configure_style(doc, "Heading 2", "TH Sarabun New", 18, bold=True)
    configure_style(doc, "Heading 3", "TH Sarabun New", 16, bold=True)

    if "Subtitle" not in doc.styles:
        subtitle = doc.styles.add_style("Subtitle", WD_STYLE_TYPE.PARAGRAPH)
        subtitle.font.name = "TH Sarabun New"
        subtitle.font.size = Pt(16)
        subtitle.font.italic = True
        subtitle.paragraph_format.space_after = Pt(6)
        subtitle.paragraph_format.line_spacing = 1.15

    if "Code Block" not in doc.styles:
        code_style = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Consolas"
        code_style.font.size = Pt(9)
        code_style.paragraph_format.left_indent = Inches(0.25)
        code_style.paragraph_format.space_after = Pt(6)
        code_style.paragraph_format.line_spacing = 1.0

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "TH Sarabun New"
        style.font.size = Pt(16)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.1

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_markdown_runs(header, title, font_name="TH Sarabun New", size_pt=10)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label = footer.add_run("หน้า ")
    set_run_font(label, "TH Sarabun New", 10)
    add_field_run(footer, "PAGE")


def normalize_markdown_text(text: str) -> str:
    text = text.replace("\u2014", "-")
    text = text.replace("\u2013", "-")
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1 (\2)", text)
    return text


def plain_text(text: str) -> str:
    text = normalize_markdown_text(text)
    for marker in ("**", "__", "*", "_", "`"):
        text = text.replace(marker, "")
    return text.strip()


def add_field_run(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(end)
    set_run_font(run, "TH Sarabun New", 10)
    return run


def add_markdown_runs(
    paragraph,
    text: str,
    *,
    font_name: str = "TH Sarabun New",
    size_pt: int = 16,
    default_bold: bool = False,
):
    text = normalize_markdown_text(text)
    parts = INLINE_PATTERN.split(text)
    for part in parts:
        if not part:
            continue

        content = part
        font = font_name
        size = size_pt
        bold = default_bold
        italic = False

        if part.startswith("**") and part.endswith("**"):
            content = part[2:-2]
            bold = True
        elif part.startswith("__") and part.endswith("__"):
            content = part[2:-2]
            bold = True
        elif part.startswith("`") and part.endswith("`"):
            content = part[1:-1]
            font = "Consolas"
            size = 10
        elif part.startswith("*") and part.endswith("*"):
            content = part[1:-1]
            italic = True
        if not content:
            continue

        run = paragraph.add_run(content)
        set_run_font(run, font, size, bold=bold)
        run.italic = italic


def is_table_line(line: str) -> bool:
    return line.strip().startswith("|") and line.strip().endswith("|")


def is_separator_row(cells: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def normalize_table_rows(rows: list[list[str]]) -> list[list[str]]:
    if not rows:
        return rows

    max_cols = max(len(row) for row in rows)
    normalized: list[list[str]] = []
    for row in rows:
        if len(row) < max_cols:
            normalized.append(row + [""] * (max_cols - len(row)))
        else:
            normalized.append(row[:max_cols])
    return normalized


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return

    rows = normalize_table_rows(rows)
    header = rows[0]
    body = rows[1:]

    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for idx, value in enumerate(header):
        para = hdr_cells[idx].paragraphs[0]
        add_markdown_runs(para, value, font_name="TH Sarabun New", size_pt=16, default_bold=True)

    for row in body:
        tr = table.add_row().cells
        for idx, value in enumerate(row):
            para = tr[idx].paragraphs[0]
            add_markdown_runs(para, value, font_name="TH Sarabun New", size_pt=16)

    doc.add_paragraph("")


def add_code_block(doc: Document, text: str):
    para = doc.add_paragraph(style="Code Block")
    run = para.add_run(text.rstrip())
    set_run_font(run, "Consolas", 9)


def add_body_paragraph(doc: Document, text: str):
    para = doc.add_paragraph()
    add_markdown_runs(para, text)
    return para


def extract_front_matter(lines: list[str], fallback_title: str) -> tuple[str, str | None, int]:
    idx = 0
    while idx < len(lines) and not lines[idx].strip():
        idx += 1

    title = fallback_title
    metadata = None

    if idx < len(lines) and lines[idx].strip().startswith("# "):
        title = plain_text(lines[idx].strip()[2:])
        idx += 1

    while idx < len(lines) and not lines[idx].strip():
        idx += 1

    if idx < len(lines):
        stripped = lines[idx].strip()
        if stripped.startswith("**") and stripped.endswith("**"):
            metadata = plain_text(stripped)
            idx += 1

    while idx < len(lines) and not lines[idx].strip():
        idx += 1

    if idx < len(lines) and HORIZONTAL_RULE_PATTERN.fullmatch(lines[idx].strip()):
        idx += 1

    while idx < len(lines) and not lines[idx].strip():
        idx += 1

    return title, metadata, idx


def add_cover_page(doc: Document, title: str, metadata: str | None, source_name: str):
    title_para = doc.add_paragraph(style="Title")
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_markdown_runs(title_para, title, font_name="TH Sarabun New", size_pt=26, default_bold=True)

    if metadata:
        subtitle_para = doc.add_paragraph(style="Subtitle")
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_markdown_runs(subtitle_para, metadata, font_name="TH Sarabun New", size_pt=16)

    doc.add_paragraph("")

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_markdown_runs(
        info_para,
        f"Human-readable edition generated from `doc/ai/{source_name}`",
        font_name="TH Sarabun New",
        size_pt=14,
    )

    generated_para = doc.add_paragraph()
    generated_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_markdown_runs(generated_para, f"Generated on: {GENERATED_ON}", font_name="TH Sarabun New", size_pt=14)

    note_para = doc.add_paragraph()
    note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_markdown_runs(
        note_para,
        "เปิดไฟล์ใน Microsoft Word แล้วกด Update Field หากต้องการ refresh สารบัญและเลขหน้า",
        font_name="TH Sarabun New",
        size_pt=13,
    )

    doc.add_page_break()

    toc_heading = doc.add_paragraph(style="Heading 1")
    add_markdown_runs(toc_heading, "สารบัญ", font_name="TH Sarabun New", size_pt=20, default_bold=True)

    toc_para = doc.add_paragraph()
    add_field_run(toc_para, r'TOC \o "1-3" \h \z \u')

    doc.add_page_break()


def render_markdown(md_path: Path, output_path: Path) -> tuple[str, str | None]:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    title, metadata, start_idx = extract_front_matter(lines, md_path.stem)

    doc = Document()
    configure_document(doc, title)
    add_cover_page(doc, title, metadata, md_path.name)

    in_code_block = False
    code_buffer: list[str] = []
    table_buffer: list[str] = []
    last_blank = False

    def flush_code():
        nonlocal code_buffer, last_blank
        if not code_buffer:
            return
        add_code_block(doc, "\n".join(code_buffer))
        code_buffer = []
        last_blank = False

    def flush_table():
        nonlocal table_buffer, last_blank
        if not table_buffer:
            return
        rows = [split_table_row(line) for line in table_buffer]
        if len(rows) > 1 and is_separator_row(rows[1]):
            rows.pop(1)
        add_table(doc, rows)
        table_buffer = []
        last_blank = False

    for raw_line in lines[start_idx:]:
        line = raw_line.rstrip()

        if line.startswith("```"):
            if in_code_block:
                flush_code()
            flush_table()
            in_code_block = not in_code_block
            continue

        if in_code_block:
            code_buffer.append(line)
            continue

        if is_table_line(line):
            table_buffer.append(line)
            continue

        flush_table()
        stripped = line.strip()

        if not stripped or HORIZONTAL_RULE_PATTERN.fullmatch(stripped):
            if not last_blank:
                doc.add_paragraph("")
                last_blank = True
            continue

        if stripped.startswith("# "):
            para = doc.add_paragraph(style="Heading 1")
            add_markdown_runs(para, stripped[2:], font_name="TH Sarabun New", size_pt=20, default_bold=True)
            last_blank = False
            continue

        if stripped.startswith("## "):
            para = doc.add_paragraph(style="Heading 2")
            add_markdown_runs(para, stripped[3:], font_name="TH Sarabun New", size_pt=18, default_bold=True)
            last_blank = False
            continue

        if stripped.startswith("### "):
            para = doc.add_paragraph(style="Heading 3")
            add_markdown_runs(para, stripped[4:], font_name="TH Sarabun New", size_pt=16, default_bold=True)
            last_blank = False
            continue

        if stripped.startswith("#### "):
            para = doc.add_paragraph(style="Heading 3")
            add_markdown_runs(para, stripped[5:], font_name="TH Sarabun New", size_pt=16, default_bold=True)
            last_blank = False
            continue

        if stripped == ">":
            if not last_blank:
                doc.add_paragraph("")
                last_blank = True
            continue

        if stripped.startswith(">"):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25)
            quote_text = stripped[1:].lstrip()
            if quote_text:
                add_markdown_runs(para, quote_text, font_name="TH Sarabun New", size_pt=16)
            last_blank = False
            continue

        if re.match(r"^[-*]\s+", stripped):
            para = doc.add_paragraph(style="List Bullet")
            add_markdown_runs(para, re.sub(r"^[-*]\s+", "", stripped), font_name="TH Sarabun New", size_pt=16)
            last_blank = False
            continue

        if re.match(r"^\d+\.\s+", stripped):
            para = doc.add_paragraph(style="List Number")
            add_markdown_runs(para, re.sub(r"^\d+\.\s+", "", stripped), font_name="TH Sarabun New", size_pt=16)
            last_blank = False
            continue

        add_body_paragraph(doc, stripped)
        last_blank = False

    flush_table()
    flush_code()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return title, metadata


def get_document_order(readme_name: str) -> list[str]:
    readme_path = AI_DIR / readme_name
    if not readme_path.exists():
        return []

    ordered_names: list[str] = []
    for line in readme_path.read_text(encoding="utf-8").splitlines():
        match = ORDER_LINE_PATTERN.match(line.strip())
        if match:
            ordered_names.append(match.group("name"))
    return ordered_names


def build_human_index(doc_infos: list[dict[str, str]], output_path: Path):
    doc = Document()
    configure_document(doc, "SB-00 - Human Document Index")

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_markdown_runs(title, "SB-00 - Human Document Index", font_name="TH Sarabun New", size_pt=26, default_bold=True)

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_markdown_runs(
        subtitle,
        "สารบัญรวมของไฟล์ Word สำหรับคนอ่าน",
        font_name="TH Sarabun New",
        size_pt=16,
    )

    intro = doc.add_paragraph()
    add_markdown_runs(
        intro,
        "ไฟล์นี้ช่วยนำทางชุดเอกสาร human-readable ที่สร้างจาก `doc/ai/` และควรใช้คู่กับไฟล์ Word ในโฟลเดอร์เดียวกัน",
        font_name="TH Sarabun New",
        size_pt=16,
    )

    guide_heading = doc.add_paragraph(style="Heading 1")
    add_markdown_runs(guide_heading, "How To Read", font_name="TH Sarabun New", size_pt=20, default_bold=True)

    for item in [
        "เริ่มจาก Master Assumptions เพื่อดู baseline กลาง",
        "ใช้ Decision Register เพื่อดู final choice ของ decision สำคัญที่ล็อกแล้ว",
        "อ่าน Dev Plan และ Business Roadmap ก่อนลงรายละเอียดเชิงเทคนิค",
        "เปิด Backend, Firmware, และ Procurement ตามหัวข้อที่ต้องการลงลึก",
    ]:
        para = doc.add_paragraph(style="List Bullet")
        add_markdown_runs(para, item, font_name="TH Sarabun New", size_pt=16)

    table_heading = doc.add_paragraph(style="Heading 1")
    add_markdown_runs(table_heading, "Document List", font_name="TH Sarabun New", size_pt=20, default_bold=True)

    rows = [["Order", "Document", "Version / Notes", "Word File"]]
    for index, info in enumerate(doc_infos, start=1):
        rows.append(
            [
                str(index),
                info["title"],
                info.get("metadata", "") or "-",
                info["human_name"],
            ]
        )
    add_table(doc, rows)

    closing = doc.add_paragraph()
    add_markdown_runs(
        closing,
        f"Generated on: {GENERATED_ON}",
        font_name="TH Sarabun New",
        size_pt=14,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main():
    HUMAN_DIR.mkdir(parents=True, exist_ok=True)

    ordered_names = get_document_order("HUMAN_README.md")
    available_files = {path.name: path for path in AI_DIR.glob("*.md") if path.name not in SKIP_FILES}

    ordered_paths: list[Path] = []
    seen: set[str] = set()
    for name in ordered_names:
        path = available_files.get(name)
        if path:
            ordered_paths.append(path)
            seen.add(name)

    if not ordered_paths:
        ordered_names = get_document_order("README.md")
        for name in ordered_names:
            path = available_files.get(name)
            if path and name not in seen:
                ordered_paths.append(path)
                seen.add(name)

        for name in sorted(available_files):
            if name not in seen:
                ordered_paths.append(available_files[name])

    doc_infos: list[dict[str, str]] = []
    generated_names: set[str] = set()
    for md_path in ordered_paths:
        output_path = HUMAN_DIR / f"{md_path.stem}.docx"
        title, metadata = render_markdown(md_path, output_path)
        generated_names.add(output_path.name)
        doc_infos.append(
            {
                "source_name": md_path.name,
                "title": title,
                "metadata": metadata or "",
                "human_name": output_path.name,
            }
        )
        print(f"generated {output_path}")

    index_path = HUMAN_DIR / "SB-00_Document_Index_v1_1.docx"
    build_human_index(doc_infos, index_path)
    print(f"generated {index_path}")
    generated_names.add(index_path.name)

    for existing in HUMAN_DIR.glob("*.docx"):
        if existing.name.startswith("~$"):
            continue
        if existing.name not in generated_names:
            existing.unlink()


if __name__ == "__main__":
    main()
